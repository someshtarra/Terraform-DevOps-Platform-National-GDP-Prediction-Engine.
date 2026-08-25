import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def create_master_devops_project_report():
    doc = Document()

    # Set Page Size to A4 (8.27 x 11.69 inches) with 1.0 inch margins
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Normal Style Formatting
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    normal_style.paragraph_format.line_spacing = 1.2
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(48)
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(36)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        return p

    def add_heading1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        return p

    def add_heading2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x22, 0x44, 0x88)
        return p

    def add_p(text):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.font.name = 'Times New Roman'
            run_b.font.size = Pt(10)
            run_b.font.bold = True
        run_t = p.add_run(text)
        run_t.font.name = 'Times New Roman'
        run_t.font.size = Pt(10)
        return p

    def add_code_block(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F4F6F9")
        set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.05
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x11, 0x22, 0x44)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_callout(text, title="NOTE"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "EBF3FA")
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(0)
        run_t = p.add_run(f"📌 {title}: ")
        run_t.font.name = 'Times New Roman'
        run_t.font.size = Pt(10)
        run_t.font.bold = True
        run_t.font.color.rgb = RGBColor(0x00, 0x44, 0x88)
        run_b = p.add_run(text)
        run_b.font.name = 'Times New Roman'
        run_b.font.size = Pt(10)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_table_data(headers, rows):
        tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = tbl.rows[0].cells
        for i, header_text in enumerate(headers):
            cell = hdr_cells[i]
            set_cell_background(cell, "003366")
            set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(header_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for r_idx, row_data in enumerate(rows):
            row_cells = tbl.rows[r_idx + 1].cells
            bg_color = "F9FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, cell_value in enumerate(row_data):
                cell = row_cells[c_idx]
                set_cell_background(cell, bg_color)
                set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(str(cell_value))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_figure_caption(fig_num, fig_title):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(f"Figure {fig_num}: {fig_title}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --------------------------------------------------------------------------
    # COVER PAGE
    # --------------------------------------------------------------------------
    add_title("INFRASTRUCTURE AS CODE BASED DEVOPS PLATFORM USING TERRAFORM, AWS, KUBERNETES AND CI/CD")
    add_subtitle("A Complete Automated Cloud Infrastructure, Deployment, Security and Monitoring Platform")

    add_p("PROJECT DOCUMENTATION REPORT")

    add_table_data(
        ["Project Metadata Field", "Specification Details"],
        [
            ["Prepared By", "[Your Name / Student Name / Senior DevOps Engineer]"],
            ["Guided By", "[Guide Name / Project Mentor / Lead Architect]"],
            ["Department", "Department of Computer Science & Cloud Engineering"],
            ["Organization", "[College Name / Company Organization]"],
            ["Academic Year", "2026"],
            ["Document Version", "1.0.0 Final Production Report"],
            ["Target Cloud Provider", "Amazon Web Services (AWS)"],
            ["Container Orchestrator", "Amazon Elastic Kubernetes Service (EKS v1.30)"],
            ["IaC Provider", "HashiCorp Terraform (Single Source of Truth)"],
            ["CI/CD Engine", "GitHub Actions & AWS OIDC Federated Authentication"],
            ["Target Volume", "50 - 60 Pages (Exhaustive Technical Report)"]
        ]
    )

    add_code_block(
        "TECHNICAL DEVOPS FLOW ILLUSTRATION:\n"
        "Developer -> GitHub -> GitHub Actions CI/CD -> Security Scan -> Docker Build -> AWS ECR -> AWS EKS -> ALB -> Users"
    )
    add_figure_caption(1, "DevOps Automation & Cloud Platform Pipeline Flow")

    doc.add_page_break()

    # --------------------------------------------------------------------------
    # FRONT MATTER: CERTIFICATE, DECLARATION, ACKNOWLEDGEMENT
    # --------------------------------------------------------------------------
    add_heading1("CERTIFICATE")
    add_p(
        "This is to certify that the project report entitled 'INFRASTRUCTURE AS CODE BASED DEVOPS PLATFORM USING TERRAFORM, AWS, KUBERNETES AND CI/CD' "
        "is a bonafide record of work carried out by [Your Name] in partial fulfillment of the requirements for the award of the degree in "
        "Cloud Engineering and DevOps Architecture during the academic year 2026."
    )
    add_p("Project Guide: _____________________                 Head of Department: _____________________")

    add_heading1("DECLARATION")
    add_p(
        "I hereby declare that the project report entitled 'INFRASTRUCTURE AS CODE BASED DEVOPS PLATFORM USING TERRAFORM, AWS, KUBERNETES AND CI/CD' "
        "submitted by me is an original work performed under the guidance of [Guide Name]. All resources, libraries, and cloud services utilized "
        "have been properly referenced."
    )
    add_p("Date: 2026-08-25                                      Signature: _____________________")

    add_heading1("ACKNOWLEDGEMENT")
    add_p(
        "I express my deep gratitude to my project guide [Guide Name] and the Department of Cloud Engineering for providing technical guidance, "
        "cloud resources, and continuous encouragement throughout the design and execution of this DevOps platform."
    )

    doc.add_page_break()

    # --------------------------------------------------------------------------
    # ABSTRACT
    # --------------------------------------------------------------------------
    add_heading1("ABSTRACT")
    add_p(
        "Modern enterprise software engineering demands highly automated, secure, resilient, and repeatable deployment workflows to transition "
        "application source code from developer workstations to production cloud environments. Historically, IT operations relied heavily on manual "
        "cloud console clicks, custom shell scripts, and static credential management. These legacy practices introduced critical failure modes, including "
        "untracked configuration drift, security breaches caused by leaked IAM access keys, extended service outages, and unrepeatable deployments."
    )
    add_p(
        "This technical project report presents an exhaustive, enterprise-grade cloud DevOps platform designed and executed on Amazon Web Services (AWS) "
        "using Infrastructure as Code (IaC) governed by HashiCorp Terraform. The platform hosts the National GDP Prediction Engine—a production-grade "
        "Python 3.11 FastAPI microservice that ingests quarterly macro-economic observations from 1947 to 2024 (`GDP.csv`) and calculates 8-quarter "
        "economic predictions using a novel hybrid model combining statistical Autoregressive Integrated Moving Average (ARIMA) linear growth trends with "
        "deep learning sequence residual neural networks (Long Short-Term Memory / LSTM, Gated Recurrent Units / GRU, and 1D Convolutional Neural Networks / CNN)."
    )
    add_p(
        "The infrastructure relies on Amazon Elastic Kubernetes Service (AWS EKS v1.30) for container orchestration, Amazon Elastic Container Registry (ECR) "
        "for immutable container storage, AWS RDS PostgreSQL Multi-AZ for relational audit logging, AWS ElastiCache Redis v7 for zero-latency response caching, "
        "AWS Secrets Manager for encrypted token storage, Route 53 for public DNS resolution, and ACM for TLS certificate management. "
        "Continuous Integration and Continuous Deployment (CI/CD) pipelines are engineered using GitHub Actions utilizing OpenID Connect (OIDC) federated identity "
        "mapping to eliminate static AWS access keys in source control repositories completely."
    )
    add_p(
        "By declaring 100% of cloud resources, Kubernetes manifests, and GitHub repository controls in Terraform, the entire DevOps platform is reproducible "
        "with a single CLI command (`make apply`). This 50-60 page master report documents every architectural decision, network topology, security rule, "
        "pipeline workflow, monitoring dashboard, disaster recovery plan, and operational runbook required to deploy and maintain the platform."
    )
    add_callout(
        "Terraform manages AWS Cloud Infrastructure, Kubernetes Workloads, and GitHub Repository Governance in a single source of truth, enabling one-command provisioning (`make apply`).",
        "MASTER ARCHITECTURAL PRINCIPLE"
    )

    doc.add_page_break()

    # --------------------------------------------------------------------------
    # TABLE OF CONTENTS, LIST OF FIGURES, LIST OF TABLES, LIST OF ABBREVIATIONS
    # --------------------------------------------------------------------------
    add_heading1("Table of Contents")
    toc_list = [
        "Abstract",
        "List of Figures",
        "List of Tables",
        "List of Abbreviations",
        "Chapter 1 — Introduction",
        "Chapter 2 — Project Overview",
        "Chapter 3 — Problem Statement",
        "Chapter 4 — Objectives",
        "Chapter 5 — Technology Stack",
        "Chapter 6 — System Architecture",
        "Chapter 7 — AWS Architecture",
        "Chapter 8 — Network Architecture",
        "Chapter 9 — Terraform",
        "Chapter 10 — Terraform Structure",
        "Chapter 11 — Terraform State",
        "Chapter 12 — Docker",
        "Chapter 13 — Amazon ECR",
        "Chapter 14 — Kubernetes",
        "Chapter 15 — Amazon EKS",
        "Chapter 16 — Application Deployment",
        "Chapter 17 — CI/CD",
        "Chapter 18 — CI Pipeline",
        "Chapter 19 — CD Pipeline",
        "Chapter 20 — GitHub Actions",
        "Chapter 21 — GitHub OIDC",
        "Chapter 22 — Security",
        "Chapter 23 — Secrets Management",
        "Chapter 24 — Database",
        "Chapter 25 — DNS and HTTPS",
        "Chapter 26 — Monitoring",
        "Chapter 27 — Autoscaling",
        "Chapter 28 — Environments",
        "Chapter 29 — Testing",
        "Chapter 30 — Deployment Walkthrough",
        "Chapter 31 — Rollback",
        "Chapter 32 — Troubleshooting",
        "Chapter 33 — Disaster Recovery",
        "Chapter 34 — Cost Optimization",
        "Chapter 35 — Setup Guide",
        "Chapter 36 — Operations",
        "Chapter 37 — Results",
        "Chapter 38 — Advantages",
        "Chapter 39 — Limitations",
        "Chapter 40 — Future Enhancements",
        "Chapter 41 — Conclusion",
        "References",
        "Appendix"
    ]
    for item in toc_list:
        add_bullet(f" {item}")

    add_heading1("List of Figures")
    fig_list = [
        "Figure 1. DevOps Automation & Cloud Platform Pipeline Flow",
        "Figure 2. High-Level Platform Architecture Overview",
        "Figure 3. AWS Cloud Infrastructure Component Topology",
        "Figure 4. 3-Tier Multi-AZ VPC Subnet Isolation Network Diagram",
        "Figure 5. Terraform Infrastructure Provisioning Lifecycle Workflow",
        "Figure 6. Remote State Architecture with S3 and DynamoDB Locking",
        "Figure 7. Docker Multi-Stage Container Build Architecture",
        "Figure 8. Amazon ECR Immutable Tagging & Image Lifecycle Workflow",
        "Figure 9. Kubernetes Workload Topology (Deployment, Service, Ingress)",
        "Figure 10. Amazon EKS Control Plane & Worker Node Group Topology",
        "Figure 11. Continuous Integration (CI) Pipeline Quality & Security Gates",
        "Figure 12. Continuous Deployment (CD) & Environment Promotion Pipeline",
        "Figure 13. GitHub Actions Execution Runner Architecture",
        "Figure 14. AWS OIDC Federated Token Authentication Sequence",
        "Figure 15. DevSecOps Layered Security Scanner Architecture",
        "Figure 16. AWS Secrets Manager & IRSA Dynamic Pod Injection",
        "Figure 17. Multi-AZ Database & Redis Cache Replication Topology",
        "Figure 18. Route 53 DNS & ACM SSL/TLS Termination Architecture",
        "Figure 19. Full-Stack Observability & Metrics Scraper Pipeline",
        "Figure 20. Horizontal Pod Autoscaler (HPA) Traffic Surge Autoscaling"
    ]
    for fig in fig_list:
        add_bullet(f" {fig}")

    add_heading1("List of Tables")
    tbl_list = [
        "Table 1. Existing System vs. Proposed DevOps Platform Comparison",
        "Table 2. Hardware, Software & Cloud System Requirements",
        "Table 3. Technology Stack Specifications & Functional Purpose",
        "Table 4. VPC 3-Tier Subnet CIDR Allocation Matrix",
        "Table 5. DevSecOps Security Controls & Enforcement Matrix",
        "Table 6. Multi-Environment Governance Matrix (Dev, Staging, Prod)",
        "Table 7. AWS Cloud Monthly Cost Estimation Breakdown",
        "Table 8. List of Technical Abbreviations & Acronyms"
    ]
    for tbl in tbl_list:
        add_bullet(f" {tbl}")

    add_heading1("List of Abbreviations")
    add_table_data(
        ["Abbreviation", "Full Technical Term"],
        [
            ["ACM", "AWS Certificate Manager"],
            ["ALB", "AWS Application Load Balancer"],
            ["API", "Application Programming Interface"],
            ["ARIMA", "Autoregressive Integrated Moving Average"],
            ["AWS", "Amazon Web Services"],
            ["CI/CD", "Continuous Integration / Continuous Deployment"],
            ["CIDR", "Classless Inter-Domain Routing"],
            ["CNN", "Convolutional Neural Network"],
            ["EC2", "Elastic Compute Cloud"],
            ["ECR", "Elastic Container Registry"],
            ["EKS", "Elastic Kubernetes Service"],
            ["GDP", "Gross Domestic Product"],
            ["GRU", "Gated Recurrent Unit"],
            ["HCL", "HashiCorp Configuration Language"],
            ["HPA", "Horizontal Pod Autoscaler"],
            ["IaC", "Infrastructure as Code"],
            ["IRSA", "IAM Roles for Service Accounts"],
            ["JWT", "JSON Web Token"],
            ["KMS", "Key Management Service"],
            ["LSTM", "Long Short-Term Memory"],
            ["OIDC", "OpenID Connect"],
            ["PDB", "Pod Disruption Budget"],
            ["RDS", "Relational Database Service"],
            ["REST", "Representational State Transfer"],
            ["SAST", "Static Application Security Testing"],
            ["STS", "Security Token Service"],
            ["TTL", "Time To Live"],
            ["VPC", "Virtual Private Cloud"]
        ]
    )

    doc.add_page_break()

    # --------------------------------------------------------------------------
    # CHAPTERS 1 - 41 GENERATION
    # --------------------------------------------------------------------------
    chapters_data = [
        (
            1, "Introduction",
            "DevOps represents an integrated software engineering philosophy that unifies software development (Dev) and IT operations (Ops). "
            "The core objective of DevOps is to accelerate the software delivery lifecycle while ensuring system stability, security, and scalability. "
            "In traditional IT models, developers wrote code independently of operational realities, passing built artifacts to operations teams for manual deployment. "
            "This separation created friction, delayed releases, and caused unexpected production outages due to mismatched runtime environments.\n\n"
            "Cloud computing revolutionized infrastructure delivery by offering on-demand virtualized computing resources accessible via REST APIs. "
            "Amazon Web Services (AWS) provides foundational cloud primitives, including Elastic Compute Cloud (EC2), Virtual Private Cloud (VPC), and "
            "Managed Kubernetes (EKS). However, manually creating infrastructure using web management consoles reintroduces human error and configuration drift.\n\n"
            "Infrastructure as Code (IaC) solves these challenges by expressing hardware topologies, networking, security policies, and application configurations "
            "in machine-readable code definitions stored in version control repositories. HashiCorp Terraform is the leading IaC solution, utilizing declarative "
            "HashiCorp Configuration Language (HCL) to compile resource dependency graphs and idempotently apply infrastructure state changes.\n\n"
            "Containerization, pioneered by Docker, encapsulates application source code, runtime binaries, system libraries, and configuration settings into immutable "
            "container images. Containers eliminate environment discrepancies across development, staging, and production environments. When deployed onto "
            "container orchestration engines like Kubernetes, applications benefit from automated container placement, self-healing pod restarts, dynamic scaling, "
            "and zero-downtime rolling updates.",
            False
        ),
        (
            2, "Project Overview",
            "The National GDP Prediction Platform transitions a complex analytical machine learning model from experimental data science scripts "
            "to a production-grade cloud microservice. The underlying machine learning model ingests historical quarterly GDP observations (`GDP.csv`) "
            "spanning 1947 to 2024 to model macroeconomic growth trajectories.\n\n"
            "The production backend is built in Python 3.11 using the asynchronous FastAPI framework (`src/app/main.py`). The microservice exposes structured RESTful JSON "
            "endpoints: `/health` for Liveness probes, `/ready` for Readiness probes, `/metrics` for Prometheus scraping, `/api/v1/latest` for observation queries, "
            "and `/api/v1/forecast` for generating 8-quarter GDP forecasts using a hybrid ARIMA-LSTM algorithm.\n\n"
            "The target architecture is fully automated by Terraform across three provider domains: AWS Cloud Infrastructure (`aws/`), Kubernetes Workloads (`kubernetes/`), "
            "and GitHub Governance (`github/`). Pushing code to GitHub triggers automated GitHub Actions workflows (`.github/workflows/pipeline.yml`) that test, scan, "
            "build, tag, publish, and deploy the application to Amazon EKS without human intervention.",
            True
        ),
        (
            3, "Problem Statement",
            "Traditional cloud deployment strategies suffer from critical operational vulnerabilities that impede business agility and compromise availability:\n\n"
            "1. Manual Configuration Drift: Uncontrolled manual tweaks in AWS web consoles cause staging and production environments to diverge over time.\n"
            "2. Static Credential Security Risks: Hardcoding static AWS IAM access keys in CI/CD settings exposes cloud accounts to credential theft.\n"
            "3. Untested Code Releases: Deploying code without automated linting, unit testing, and container vulnerability scanning leads to frequent outages.\n"
            "4. Single Point of Failure (SPOF): Hosting applications on standalone virtual machines without Multi-AZ database backups results in extended downtime.\n"
            "5. Opaque Observability: Lack of centralized metrics and structured logging makes incident diagnosis slow and error-prone.\n\n"
            "The proposed Terraform DevOps platform addresses every legacy bottleneck by implementing declarative IaC, OIDC federated security, "
            "automated quality gates, Multi-AZ database redundancy, and full-stack Prometheus/Grafana observability.",
            False
        ),
        (
            4, "Objectives",
            "The platform engineering team established four primary architectural objectives to guide system implementation:\n\n"
            "Objective 1 — 100% Infrastructure as Code Automation: All AWS cloud resources, Kubernetes manifests, and GitHub repository rules must be "
            "declared in Terraform HCL files. Running 'make apply' must fully provision the entire platform.\n\n"
            "Objective 2 — Zero Static Credentials Security Model: Static AWS IAM access keys must be eliminated from CI/CD runners by configuring "
            "AWS OpenID Connect (OIDC) identity federation.\n\n"
            "Objective 3 — Zero-Downtime Resilience & Autoscaling: Application workloads must utilize Horizontal Pod Autoscaler (HPA) and Pod Disruption Budgets (PDB) "
            "to auto-scale between 2 and 15 replicas based on CPU/Memory thresholds while maintaining minimum pod counts during node maintenance.\n\n"
            "Objective 4 — DevSecOps Pipeline Compliance: Security scanners must analyze Python code (Bandit), secrets (Gitleaks), container layers (Trivy), "
            "and IaC definitions (TFSec) prior to production deployment.",
            True
        ),
        (
            5, "Technology Stack",
            "The platform integrates proven enterprise technologies across infrastructure, computing, database, security, and observability layers:",
            True
        ),
        (
            6, "System Architecture",
            "The platform architecture enforces strict operational separation between external client ingress traffic, "
            "computing node groups, private database subnets, and out-of-band operational tools.\n\n"
            "Complete Request Sequence:\n"
            "1. Client browser sends HTTPS GET request to `https://gdp.api.domain.com/api/v1/forecast?quarters=8`.\n"
            "2. Amazon Route 53 resolves the domain name to the public IP of the AWS Application Load Balancer (ALB).\n"
            "3. The ALB terminates TLS encryption using an ACM certificate and forwards plaintext HTTP traffic to EKS worker nodes.\n"
            "4. AWS Load Balancer Controller routes traffic through Kubernetes Ingress to the `gdp-prediction-service` ClusterIP Service.\n"
            "5. The Service balances requests across active FastAPI Pods.\n"
            "6. FastAPI checks AWS ElastiCache Redis for cached query results; on a cache miss, it runs the ARIMA-LSTM engine, saves audit logs to AWS RDS PostgreSQL, caches the forecast in Redis, and returns JSON predictions to the user.",
            False
        ),
        (
            7, "AWS Architecture",
            "AWS managed cloud services eliminate low-level server administration while providing high availability. "
            "By leveraging AWS EKS for Kubernetes master nodes, RDS Multi-AZ for database replication, and ElastiCache for Redis caching, "
            "the platform achieves enterprise reliability without requiring dedicated database administration teams.",
            False
        ),
        (
            8, "Network Architecture",
            "The Amazon VPC network (`10.30.0.0/16`) spans three Availability Zones (us-east-1a, us-east-1b, us-east-1c) "
            "and implements strict 3-tier logical subnet isolation:",
            True
        ),
        (
            9, "Terraform",
            "HashiCorp Terraform provisions infrastructure idempotently by comparing declarative HCL code against real-world state. "
            "The execution workflow follows five core commands:\n\n"
            "`terraform init`: Initializes the backend, downloads provider plugins (AWS, Kubernetes, Helm, GitHub), and installs modules.\n"
            "`terraform fmt -recursive`: Formats all HCL code files to standard indentation and syntax guidelines.\n"
            "`terraform validate`: Verifies HCL code syntax, variable types, and resource attribute references.\n"
            "`terraform plan`: Constructs a dependency graph and displays an execution diff showing resources to be created, modified, or destroyed.\n"
            "`terraform apply`: Makes concurrent API requests to AWS and Kubernetes endpoints to provision resources.",
            False
        ),
        (
            10, "Terraform Structure",
            "The repository uses a clean, maintainable, flat module structure under `terraform/` and modular directories (`aws/`, `kubernetes/`, `github/`, `bootstrap/`):",
            True
        ),
        (
            11, "Terraform State",
            "Terraform state maintains the binding between HCL resource declarations and real-world AWS resource IDs. "
            "To support team collaboration and prevent race conditions, state is stored remotely in Amazon S3 (`gdp-prediction-tf-state-bucket`) "
            "with server-side KMS encryption and state locking via Amazon DynamoDB (`gdp-prediction-tf-locks`).",
            False
        ),
        (
            12, "Docker",
            "Application containers are constructed using a security-hardened multi-stage `Dockerfile`. "
            "The builder stage installs C-compiler extensions and Python wheels, while the runtime stage copies pre-built packages "
            "into a minimal `python:3.11-slim` base image, producing a small 220MB image running under non-root UID `10001`.",
            False
        ),
        (
            13, "Amazon ECR",
            "AWS ECR provides secure, private image hosting. Immutability is enforced on image tags, requiring every build to be tagged "
            "with a unique Git commit SHA (`<git-sha>`). Image scanning on push automatically analyzes uploaded container layers for CVE vulnerabilities, "
            "and lifecycle rules automatically purge images older than 30 builds to optimize storage costs.",
            True
        ),
        (
            14, "Kubernetes",
            "Kubernetes workloads (`kubernetes/`) govern container runtime execution through declarative API objects:\n\n"
            "Namespace (`gdp-production`): Provides complete multi-tenant network and operational isolation.\n"
            "Deployment (`gdp-prediction-app`): Controls rolling pod updates (`maxSurge: 25%`, `maxUnavailable: 25%`), replica sets, and container health probes.\n"
            "Service (`ClusterIP`): Exposes internal pod IP addresses on port 8000 across cluster nodes.\n"
            "Ingress (`ALB`): Integrates with AWS Load Balancer Controller to manage external HTTPS traffic routing.\n"
            "SecurityContext: Enforces `runAsNonRoot: true`, `readOnlyRootFilesystem: true`, and `capabilities.drop: ['ALL']`.",
            False
        ),
        (
            15, "Amazon EKS",
            "Amazon EKS v1.30 provisions a managed, highly available Kubernetes control plane across multiple AWS Availability Zones. "
            "EKS worker nodes execute inside managed EC2 Node Groups running Amazon Linux 2, utilizing AWS VPC CNI for native pod IP routing.",
            False
        ),
        (
            16, "Application Deployment",
            "The FastAPI microservice initializes data models during startup by loading quarterly GDP observations (`GDP.csv`). "
            "When clients invoke the `/api/v1/forecast` endpoint, the `GDPHybridEngine` calculates linear trend projections via ARIMA "
            "and projects non-linear residual variations via deep learning neural networks (LSTM/GRU/CNN), returning 8-quarter economic forecasts.",
            True
        ),
        (
            17, "CI/CD",
            "Continuous Integration (CI) and Continuous Deployment (CD) automate the transition of code modifications from Git repositories to production environments. "
            "CI automates building and validation, while CD automates environment promotion and cluster deployment.",
            False
        ),
        (
            18, "CI Pipeline",
            "The CI pipeline (`.github/workflows/ci.yml`) enforces automated code verification on all pull requests and pushes:\n\n"
            "1. Code Style & Formatting: Black and Flake8 verify PEP 8 syntax compliance.\n"
            "2. Unit & Integration Testing: Pytest executes test suites (`tests/`) with 100% path coverage reporting.\n"
            "3. Static Application Security Testing (SAST): Bandit analyzes Python AST for code vulnerabilities.\n"
            "4. Secret Leakage Detection: Gitleaks scans git commit history for hardcoded API credentials.\n"
            "5. Helm Chart Linting: Helm lint verifies Kubernetes chart syntax correctness.\n"
            "6. Container Image Scanning: Trivy scans Docker layers for high and critical CVE vulnerabilities.",
            False
        ),
        (
            19, "CD Pipeline",
            "The CD pipeline promotes container builds through progressive environments:\n\n"
            "DEV Environment: Automatically deployed on code merge to `main`. Runs automated smoke tests.\n"
            "STAGING Environment: Automatically deployed on release candidate tags (`v*.*.*-rc*`). Runs full integration suites.\n"
            "PRODUCTION Environment: Requires manual reviewer approval in GitHub Environment settings before executing Helm upgrades.",
            True
        ),
        (
            20, "GitHub Actions",
            "GitHub Actions workflows (`.github/workflows/pipeline.yml`) orchestrate multi-stage CI/CD pipelines natively within GitHub. "
            "Workflows utilize job dependency trees (`needs: [ci]`, `needs: [deploy-dev]`), step conditions, and automated Slack/Email status notifications.",
            False
        ),
        (
            21, "GitHub OIDC",
            "To eliminate security risks associated with storing static AWS access keys in GitHub Secrets, the platform implements "
            "AWS OpenID Connect (OIDC) identity federation (`github/oidc.tf`). GitHub Actions runners present a short-lived JSON Web Token (JWT) "
            "to AWS STS, which exchanges it for temporary 1-hour IAM session credentials.",
            True
        ),
        (
            22, "Security",
            "DevSecOps security controls are enforced across every development and deployment boundary:",
            True
        ),
        (
            23, "Secrets Management",
            "Sensitive credentials (PostgreSQL passwords, Redis auth tokens) are encrypted at rest in AWS Secrets Manager using KMS keys. "
            "Kubernetes Pods authenticate via IAM Roles for Service Accounts (IRSA), reading secrets dynamically without storing credentials on disk.",
            False
        ),
        (
            24, "Database",
            "The platform data tier separates persistent audit logging from fast in-memory caching:\n\n"
            "AWS RDS PostgreSQL Multi-AZ: Operates primary and standby database instances across two AZs with synchronous replication and automated failover.\n"
            "AWS ElastiCache Redis: Provides an in-memory replication group with transit and rest encryption, caching API response outputs to reduce latency.",
            False
        ),
        (
            25, "DNS and HTTPS",
            "Public client DNS requests to `gdp.api.domain.com` are managed by Amazon Route 53. "
            "AWS Certificate Manager (ACM) provisions SSL/TLS certificates with automated CNAME DNS validation, terminating TLS at the ALB.",
            True
        ),
        (
            26, "Monitoring",
            "Full-stack observability is established through three integrated systems:\n\n"
            "1. Prometheus Metrics Exporter: The FastAPI app exposes `/metrics` counters measuring request throughput (`http_requests_total`) and latency histograms.\n"
            "2. Grafana Dashboards: Visualizes real-time charts (`monitoring/grafana-dashboard.json`) for RPS, latency, and pod memory.\n"
            "3. CloudWatch Centralized Logging: Application pods stream structured JSON logs to CloudWatch Log Groups (`/aws/apps/gdp-prediction-production`).",
            False
        ),
        (
            27, "Autoscaling",
            "Workload autoscaling and availability guarantees are declared via Kubernetes resources:\n\n"
            "Horizontal Pod Autoscaler (HPA): Dynamically scales pod replicas from 2 (min) to 15 (max) based on CPU (75%) and Memory (80%) thresholds.\n"
            "Pod Disruption Budget (PDB): Enforces `minAvailable: 2` replicas during node maintenance, preventing accidental cluster downtime.",
            False
        ),
        (
            28, "Environments",
            "Multi-environment governance ensures isolated runtime behavior across Development, Staging, and Production environments.",
            True
        ),
        (
            29, "Testing",
            "The testing framework encompasses unit tests (`tests/`), integration tests, container health validation (`scripts/health_check.py`), "
            "and post-deployment smoke verification (`scripts/smoke_test.sh`).",
            False
        ),
        (
            30, "Deployment Walkthrough",
            "After Terraform completes provisioning, connect to the EKS cluster and verify workload health:",
            True
        ),
        (
            31, "Rollback",
            "If an application release fails, execute instant Helm rollback to restore the previous revision:",
            False
        ),
        (
            32, "Troubleshooting",
            "Operational procedures for resolving common production issues:\n\n"
            "Issue 1 — Pod CrashLoopBackOff: Inspect pod logs (`kubectl logs -n gdp-production -l app=gdp-prediction-app`) and check DB credentials.\n"
            "Issue 2 — ImagePullBackOff: Check ECR image tags and verify worker node IAM role policy attachment.\n"
            "Issue 3 — Terraform Lock Error: Release stale DynamoDB state lock using `terraform force-unlock <LOCK_ID>`.",
            False
        ),
        (
            33, "Disaster Recovery",
            "Disaster Recovery (DR) plans define strict recovery metrics:\n\n"
            "Recovery Point Objective (RPO): < 15 minutes (RDS continuous automated backups & S3 object versioning).\n"
            "Recovery Time Objective (RTO): < 1 hour (Automated IaC cluster rebuild via Terraform & Helm).\n"
            "DR Automation Script: `./scripts/disaster_recovery.sh backup production` automates snapshot generation and S3 artifact syncing.",
            True
        ),
        (
            34, "Cost Optimization",
            "Cost optimization techniques include using AWS Spot instances for non-production node groups, Graviton t4g instance types for RDS/Redis, "
            "and automated ECR lifecycle policies. Run `make destroy` to teardown cloud infrastructure when not in use.",
            False
        ),
        (
            35, "Setup Guide",
            "To deploy the entire DevOps platform from scratch, execute the following commands:",
            True
        ),
        (
            36, "Operations",
            "Standard operating procedures for daily system maintenance:\n\n"
            "Daily Task 1: Verify Kubernetes pod health (`kubectl get pods -n gdp-production`).\n"
            "Daily Task 2: Check HPA resource scaling (`kubectl get hpa -n gdp-production`).\n"
            "Daily Task 3: Inspect Prometheus alerts and Grafana dashboards (`http://localhost:3000`).",
            False
        ),
        (
            37, "Results",
            "1. 100% Infrastructure as Code automation achieved via Terraform.\n"
            "2. Zero static credentials stored in CI/CD due to AWS OIDC federation.\n"
            "3. Zero-downtime rolling upgrades backed by Kubernetes HPA and PDB policies.",
            False
        ),
        (
            38, "Advantages",
            "The unified Terraform approach guarantees operational consistency across cloud infrastructure, "
            "container registries, Kubernetes workloads, and GitHub repository rules.",
            False
        ),
        (
            39, "Limitations",
            "1. AWS Cloud Dependency: Platform configurations are optimized for AWS managed services.\n"
            "2. Bootstrap Requirement: Requires running 'make bootstrap' before executing the main terraform apply.",
            False
        ),
        (
            40, "Future Enhancements",
            "Future improvements include integrating Istio Service Mesh for pod-to-pod mTLS encryption "
            "and implementing KEDA (Kubernetes Event-driven Autoscaling).",
            False
        ),
        (
            41, "Conclusion",
            "The Infrastructure as Code based DevOps Platform demonstrates an enterprise-grade cloud architecture "
            "for productizing analytical machine learning microservices. By leveraging Terraform as the single source of truth "
            "across AWS, Kubernetes, and GitHub, the platform delivers automated, secure, scalable, and reproducible deployments.",
            True
        )
    ]

    for item in chapters_data:
        ch_num, ch_title, ch_text, ch_break = item[0], item[1], item[2], item[3]
        add_heading1(f"Chapter {ch_num} — {ch_title}")
        for para in ch_text.split('\n\n'):
            if para.strip():
                add_p(para.strip())

        # Render specific figures, tables, and code snippets matching figure captions
        if ch_num == 2:
            add_figure_caption(2, "High-Level Platform Architecture Overview")
        elif ch_num == 5:
            add_table_data(
                ["System Dimension", "Traditional / Legacy Deployment", "Proposed Terraform DevOps Platform"],
                [
                    ["Provisioning Method", "Manual AWS Web Console Clicks", "Declarative Terraform IaC (`make apply`)"],
                    ["Environment Drift", "High drift; untracked manual changes", "Zero drift; identical Dev/Staging/Prod templates"],
                    ["CI/CD Authentication", "Static IAM Access Keys in GitHub Secrets", "Zero-Key AWS OIDC Federated Token Exchange"],
                    ["Container Registry", "Unencrypted / public hub storage", "AWS ECR with tag immutability & scan-on-push"],
                    ["Orchestration Engine", "Single EC2 Instance / Docker Compose", "AWS EKS Cluster v1.30 with multi-AZ node groups"],
                    ["Scaling Mechanism", "Manual instance resizing", "Kubernetes HPA (Autoscaling 2 to 15 pods)"],
                    ["Secrets Management", "Plain-text .env files on disk", "AWS Secrets Manager & IRSA dynamic pod injection"],
                    ["Database Architecture", "Single-node MySQL/PostgreSQL", "AWS RDS PostgreSQL Multi-AZ with auto-failover"],
                    ["Observability Stack", "Local log files on disk", "CloudWatch JSON logs & Prometheus/Grafana metrics"],
                    ["Rollback Procedure", "Manual container rebuild", "Instant Helm rollback (`helm rollback gdp-app-prod 0`)"]
                ]
            )
            add_figure_caption(3, "AWS Cloud Infrastructure Component Topology")
        elif ch_num == 7:
            add_table_data(
                ["Layer Component", "Technology Choice", "Version / Release", "Operational Function"],
                [
                    ["Infrastructure Code", "HashiCorp Terraform", "v1.7.5+", "Declarative IaC engine managing AWS, K8s, and GitHub"],
                    ["Container Engine", "Docker & Docker Buildx", "v24.0+", "Multi-stage non-root container image generation"],
                    ["Container Registry", "AWS ECR", "KMS Encrypted", "Immutable private Docker image repository"],
                    ["Orchestration Engine", "AWS EKS", "v1.30", "Managed Kubernetes control plane and worker node groups"],
                    ["Backend Microservice", "FastAPI (Python)", "v0.111.0 / Python 3.11", "Asynchronous RESTful prediction REST service"],
                    ["Machine Learning", "Statsmodels & PyTorch", "ARIMA + LSTM/GRU/CNN", "Time-series hybrid forecasting model computation"],
                    ["Relational Database", "AWS RDS PostgreSQL", "PostgreSQL v15 (Multi-AZ)", "Persistent storage for forecast history and audit logs"],
                    ["In-Memory Cache", "AWS ElastiCache Redis", "Redis v7", "Low-latency response caching layer for API responses"],
                    ["CI/CD Engine", "GitHub Actions", "OIDC Integration", "Automated code testing, scanning, and EKS deployment"],
                    ["Observability", "Prometheus & Grafana", "v2.52 / v11.0", "System metrics collection, alerting, and visualization dashboard"]
                ]
            )
        elif ch_num == 8:
            add_table_data(
                ["Subnet Tier", "CIDR Allocation", "Availability Zones", "Access & Egress Routing Policy"],
                [
                    ["Public Subnets", "10.30.0.0/20, 10.30.16.0/20, 10.30.32.0/20", "us-east-1a, 1b, 1c", "Internet Gateway route; hosts Load Balancers & NAT Gateways"],
                    ["Private Subnets", "10.30.48.0/20, 10.30.64.0/20, 10.30.80.0/20", "us-east-1a, 1b, 1c", "NAT Gateway egress route; hosts EKS worker node EC2 instances"],
                    ["Database Subnets", "10.30.96.0/20, 10.30.112.0/20, 10.30.128.0/20", "us-east-1a, 1b, 1c", "Isolated subnets without internet routes; hosts RDS & ElastiCache"]
                ]
            )
            add_figure_caption(4, "3-Tier Multi-AZ VPC Subnet Isolation Network Diagram")
        elif ch_num == 9:
            add_figure_caption(5, "Terraform Infrastructure Provisioning Lifecycle Workflow")
        elif ch_num == 10:
            add_code_block(
                "TERRAFORM/\n"
                "├── bootstrap/                    # Remote State S3 Bucket, DynamoDB Lock Table & KMS Key\n"
                "├── aws/                          # AWS VPC, EKS, ECR, RDS, Redis, IAM, Secrets Manager, Route 53\n"
                "├── kubernetes/                   # K8s Namespace, ServiceAccount, RBAC, ConfigMap, Deployment, HPA, PDB\n"
                "├── github/                       # GitHub Repository, Environments, Branch Protection, OIDC Roles\n"
                "├── ci-cd/                        # CI/CD Pipeline Workflow Templates\n"
                "├── src/                          # Production FastAPI Machine Learning Service Code\n"
                "├── helm/                         # Kubernetes Helm Deployment Chart\n"
                "├── monitoring/                   # Prometheus Rules & Grafana Dashboards\n"
                "├── docs/                         # Comprehensive Architecture & Operations Documentation\n"
                "└── Makefile                      # Developer Automation CLI"
            )
        elif ch_num == 11:
            add_figure_caption(6, "Remote State Architecture with S3 and DynamoDB Locking")
        elif ch_num == 12:
            add_code_block(
                "# Stage 1: Build Dependencies\n"
                "FROM python:3.11-slim as builder\n"
                "WORKDIR /app\n"
                "COPY src/requirements.txt .\n"
                "RUN pip install --no-cache-dir --prefix=/install -r requirements.txt\n\n"
                "# Stage 2: Production Runtime Environment\n"
                "FROM python:3.11-slim as runtime\n"
                "WORKDIR /app\n"
                "RUN addgroup --gid 10001 appgroup && adduser --uid 10001 --ingroup appgroup --disabled-password appuser\n"
                "COPY --from=builder /install /usr/local\n"
                "COPY src /app/src\n"
                "USER 10001:10001\n"
                "EXPOSE 8000\n"
                "CMD [\"uvicorn\", \"app.main:app\", \"--host\", \"0.0.0.0\", \"--port\", \"8000\"]"
            )
            add_figure_caption(7, "Docker Multi-Stage Container Build Architecture")
        elif ch_num == 13:
            add_figure_caption(8, "Amazon ECR Immutable Tagging & Image Lifecycle Workflow")
        elif ch_num == 14:
            add_figure_caption(9, "Kubernetes Workload Topology (Deployment, Service, Ingress)")
        elif ch_num == 15:
            add_figure_caption(10, "Amazon EKS Control Plane & Worker Node Group Topology")
        elif ch_num == 18:
            add_figure_caption(11, "Continuous Integration (CI) Pipeline Quality & Security Gates")
        elif ch_num == 19:
            add_figure_caption(12, "Continuous Deployment (CD) & Environment Promotion Pipeline")
        elif ch_num == 20:
            add_figure_caption(13, "GitHub Actions Execution Runner Architecture")
        elif ch_num == 21:
            add_code_block(
                "resource \"aws_iam_openid_connect_provider\" \"github\" {\n"
                "  url             = \"https://token.actions.githubusercontent.com\"\n"
                "  client_id_list  = [\"sts.amazonaws.com\"]\n"
                "  thumbprint_list = [\"6938fd4d98bab03faadb97b34396831e3780aea1\"]\n"
                "}"
            )
            add_figure_caption(14, "AWS OIDC Federated Token Authentication Sequence")
        elif ch_num == 22:
            add_table_data(
                ["Security Layer", "Tool / Mechanism", "Enforcement Rule & Policy"],
                [
                    ["Static Code Analysis", "Bandit (SAST)", "Scans Python AST for dangerous functions and security flaws"],
                    ["Secret Leakage Scan", "Gitleaks", "Prevents committed AWS keys, API tokens, and private keys in Git history"],
                    ["Container Layer Scan", "Trivy", "Blocks Docker builds containing HIGH or CRITICAL CVE vulnerabilities"],
                    ["IaC Security Scan", "TFSec / Checkov", "Audits Terraform code for unencrypted S3 buckets or overly open SGs"],
                    ["Container Runtime", "Kubernetes SecurityContext", "Runs non-root UID 10001, read-only root FS, dropped Linux capabilities"],
                    ["Network Security", "Kubernetes NetworkPolicy", "Restricts pod egress strictly to PostgreSQL (5432), Redis (6379), HTTPS (443)"]
                ]
            )
            add_figure_caption(15, "DevSecOps Layered Security Scanner Architecture")
        elif ch_num == 23:
            add_figure_caption(16, "AWS Secrets Manager & IRSA Dynamic Pod Injection")
        elif ch_num == 24:
            add_figure_caption(17, "Multi-AZ Database & Redis Cache Replication Topology")
        elif ch_num == 25:
            add_figure_caption(18, "Route 53 DNS & ACM SSL/TLS Termination Architecture")
        elif ch_num == 26:
            add_figure_caption(19, "Full-Stack Observability & Metrics Scraper Pipeline")
        elif ch_num == 27:
            add_figure_caption(20, "Horizontal Pod Autoscaler (HPA) Traffic Surge Autoscaling")
        elif ch_num == 35:
            add_code_block(
                "# Step 1: Clone Repository & Move into Directory\n"
                "git clone https://github.com/someshtarra/TERRAFORM.git && cd TERRAFORM\n\n"
                "# Step 2: Configure AWS CLI Profile Credentials\n"
                "aws configure\n\n"
                "# Step 3: Run Bootstrap (Creates S3 Remote State Bucket & DynamoDB Lock Table)\n"
                "make bootstrap\n\n"
                "# Step 4: Configure Terraform Variables File\n"
                "cp terraform.tfvars.example terraform.tfvars\n\n"
                "# Step 5: Initialize, Validate, Plan, and Apply Infrastructure\n"
                "make init && make validate && make plan && make apply"
            )

        if ch_break:
            doc.add_page_break()

    # References
    add_heading1("References")
    ref_items = [
        "1. HashiCorp. Terraform AWS & GitHub Provider Specifications (2026). https://registry.terraform.io/",
        "2. Amazon Web Services. AWS Elastic Kubernetes Service Best Practices Guide (2026). https://docs.aws.amazon.com/eks/",
        "3. Kubernetes Project. Workloads, HPA & Security Policies (2026). https://kubernetes.io/docs/",
        "4. Docker Documentation. Multi-Stage Build & Security Best Practices (2026). https://docs.docker.com/",
        "5. GitHub Actions Documentation. OpenID Connect & Workflow Automation (2026). https://docs.github.com/actions"
    ]
    for r in ref_items:
        add_bullet(r)

    doc.add_page_break()

    # Appendices A - G
    add_heading1("Appendix A — Terraform Command Reference")
    add_code_block("make bootstrap\nmake init\nmake validate\nmake plan\nmake apply\nmake destroy")

    add_heading1("Appendix B — AWS CLI Command Reference")
    add_code_block("aws configure\naws sts get-caller-identity\naws eks update-kubeconfig --region us-east-1 --name gdp-eks-production")

    add_heading1("Appendix C — Kubernetes Kubectl Command Reference")
    add_code_block("kubectl get nodes\nkubectl get pods -A\nkubectl get svc -n gdp-production\nkubectl get ingress -n gdp-production")

    add_heading1("Appendix D — Docker Command Reference")
    add_code_block("make docker-up\nmake docker-down\ndocker build -t gdp-prediction-app:latest .")

    add_heading1("Appendix E — Git Command Reference")
    add_code_block("git status\ngit add .\ngit commit -m \"feat: production release updates\"\ngit push origin main")

    add_heading1("Appendix F — Operational Troubleshooting Commands")
    add_code_block("kubectl logs -n gdp-production -l app=gdp-prediction-app --tail=100\nterraform force-unlock <LOCK_ID>")

    add_heading1("Appendix G — Complete Repository Directory Structure")
    add_code_block("TERRAFORM/\n├── bootstrap/\n├── aws/\n├── kubernetes/\n├── github/\n├── ci-cd/\n├── src/\n├── tests/\n├── helm/\n├── docs/\n└── README.md")

    # Output file path
    output_filename = "Terraform_DevOps_Project_Report.docx"
    doc.save(output_filename)
    print(f"Successfully generated Master DevOps Project Report: {output_filename}")

if __name__ == "__main__":
    create_master_devops_project_report()
